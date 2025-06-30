import streamlit as st
import pandas as pd
from pytz import timezone
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

from docx.oxml.ns import qn
import tempfile
from docx.shared import Cm 
from streamlit_option_menu import option_menu
import plotly.graph_objects as go
from plotly.colors import qualitative

MAPA_NIVEL_EVALUACION = {
    "1": "Jerárquico (1)",
    "2": "Medio (2)",
    "3": "Medio (3)",
    "4": "Medio (4)",
    "5": "Operativo (5)",
    "6": "Operativo (6)"
}

MAXIMO_PUNTAJE_FORMULARIO = {
    "1": 56,
    "2": 48,
    "3": 48,
    "4": 40,
    "5": 32,
    "6": 24
}



# ---- Vista: Evaluaciones ----
def mostrar(supabase):
    #st.header("📋 Evaluaciones realizadas")
    st.markdown("<h2 style='font-size:26px;'>📋 Evaluaciones realizadas</h2>", unsafe_allow_html=True)
    
    # Función para verificar rol activo
    def tiene_rol(*roles):
        return any(st.session_state.get("rol", {}).get(r, False) for r in roles)

    # Rol y dependencias desde sesión
    dependencia_usuario = st.session_state.get("dependencia", "")
    dependencia_general = st.session_state.get("dependencia_general", "")

    # Construir opciones de filtro de dependencia
    opciones_dependencia = []

    if tiene_rol("coordinador", "evaluador_general") and dependencia_general:
        opciones_dependencia.append(f"{dependencia_general} (todas)")

    if dependencia_usuario:
        opciones_dependencia.append(f"{dependencia_usuario} (individual)")

    dependencias_subordinadas = []
    if tiene_rol("coordinador", "evaluador_general") and dependencia_general:
        resultado = supabase.table("unidades_evaluacion")\
            .select("dependencia")\
            .eq("dependencia_general", dependencia_general)\
            .neq("dependencia", dependencia_usuario)\
            .execute()
        dependencias_subordinadas = sorted({d["dependencia"] for d in resultado.data})
        opciones_dependencia += [
            d for d in dependencias_subordinadas
            if d != dependencia_usuario and "UNIDAD RESIDUAL" not in d.upper()
        ]

    dependencia_seleccionada = st.selectbox("📂 Dependencia a visualizar:", opciones_dependencia)

    # Filtrar agentes por dependencia seleccionada
    if dependencia_seleccionada and "(todas)" in dependencia_seleccionada:
        dependencia_filtro = dependencia_general
        agentes = supabase.table("agentes").select("cuil, evaluado_2024").eq("dependencia_general", dependencia_filtro).execute().data
    elif dependencia_seleccionada and "(individual)" in dependencia_seleccionada:
        dependencia_filtro = dependencia_usuario
        agentes = supabase.table("agentes").select("cuil, evaluado_2024").eq("dependencia", dependencia_filtro).execute().data
    elif dependencia_seleccionada:
        dependencia_filtro = dependencia_seleccionada
        agentes = supabase.table("agentes").select("cuil, evaluado_2024").eq("dependencia", dependencia_filtro).execute().data
    else:
        st.warning("⚠️ Seleccione una dependencia válida para continuar.")
        return

    cuils_asignados = [a["cuil"] for a in agentes]
    total_asignados = len(cuils_asignados)
    evaluados = sum(1 for a in agentes if a.get("evaluado_2024") is True)
    porcentaje = (evaluados / total_asignados * 100) if total_asignados > 0 else 0

  #  st.markdown("# ")
    
    # Obtener evaluaciones filtradas
    evaluaciones = supabase.table("evaluaciones").select("*").in_("cuil", cuils_asignados).execute().data
    df_eval = pd.DataFrame(evaluaciones)
    
    if df_eval.empty:
        df_eval = pd.DataFrame(columns=[
            "formulario", "calificacion", "anulada", "fecha_evaluacion", "apellido_nombre",
            "puntaje_total", "evaluador", "id_evaluacion", "cuil"
        ])
    
    if "anulada" not in df_eval.columns:
        df_eval["anulada"] = False
    else:
        df_eval["anulada"] = df_eval["anulada"].fillna(False).astype(bool)
    
    if "fecha_evaluacion" in df_eval.columns and not df_eval["fecha_evaluacion"].isna().all():
        hora_arg = timezone('America/Argentina/Buenos_Aires')
        df_eval["Fecha"] = pd.to_datetime(df_eval["fecha_evaluacion"], utc=True).dt.tz_convert(hora_arg)
        df_eval["Fecha_formateada"] = df_eval["Fecha"].dt.strftime('%d/%m/%Y %H:%M')
    else:
        df_eval["Fecha_formateada"] = ""
    
    df_eval["Estado"] = df_eval["anulada"].apply(lambda x: "Anulada" if x else "Registrada")
    df_no_anuladas = df_eval[df_eval["anulada"] == False].copy()

    if "calif_puntaje" not in df_no_anuladas.columns:
        df_no_anuladas["calif_puntaje"] = df_no_anuladas.apply(
            lambda row: f"{row['calificacion']} ({row['puntaje_total']})"
            if pd.notna(row.get("calificacion")) and pd.notna(row.get("puntaje_total"))
            else "",
            axis=1
        )
    
    # Unir con datos de agentes si no están ya
    agentes_completos = supabase.table("agentes").select("*").in_("cuil", cuils_asignados).execute().data
    df_agentes = pd.DataFrame(agentes_completos)
    
    if not df_agentes.empty:
        df_no_anuladas = df_no_anuladas.merge(df_agentes[[
            "cuil", "agrupamiento", "nivel", "ingresante", "apellido_nombre"
        ]], on="cuil", how="left", suffixes=("", "_agente"))


    
    # Menú horizontal de navegación
    st.markdown(
        "<p style='font-size:14px; color:white'>🖱️ Haga clic en cada una de las opciones para visualizar los datos correspondientes.</p>",
        unsafe_allow_html=True
    )
    
    seleccion = option_menu(
        menu_title=None,
        options=["📊 INDICADORES", "✅ EVALUACIONES", "👥 AGENTES"],
 #       icons=["bar-chart-line", "clipboard-check"],
        orientation="horizontal",
        default_index=0,
        
        styles={
            "container": {
                "padding": "0!important", 
                "background-color": "transparent",
                # "max-width": "800px",  # ← Eliminar esta línea
                # "margin": "0 auto"     # ← Eliminar esta línea
            },
            "nav-link": {
                "font-size": "17px",
                "text-align": "center",
                "margin": "0 10px",
               # "flex": "1",  # ← Esto hace que se distribuyan uniformemente
                "max-width": "280px",  # ← Eliminar esta línea
                "color": "white",
                "font-weight": "bold",
                "background-color": "#F05A7E",
                "border-radius": "8px",
                "--hover-color": "#b03a3f",
            },
            "nav-link-selected": {
                "background-color": "#5EABD6",
                "color": "#0C0909",
                "font-weight": "bold",
                "border-radius": "8px",
            },
        }

    )
    
    if seleccion == "📊 INDICADORES":
        st.divider()
        st.markdown("<h2 style='font-size:20px;'>📊 Indicadores</h2>", unsafe_allow_html=True)
        cols = st.columns(3)
        with cols[0]: st.metric("👥 Total a Evaluar", total_asignados)
        with cols[1]: st.metric("✅ Evaluados", evaluados)
        with cols[2]: st.metric("📊 % Evaluación", f"{int(porcentaje)}%")                
        st.progress(min(100, int(porcentaje)), text=f"Progreso de evaluaciones registradas: {int(porcentaje)}%")

        st.markdown("---")

        st.markdown("<h2 style='font-size:20px;'>🏅 Distribución por Calificación</h2>", unsafe_allow_html=True)
        categorias = ["DESTACADO", "BUENO", "REGULAR", "DEFICIENTE"]
        calif_counts = {cat: 0 for cat in categorias}
        if not df_no_anuladas.empty and "calificacion" in df_no_anuladas.columns:
            temp_counts = df_no_anuladas["calificacion"].value_counts()
            for cat in categorias:
                calif_counts[cat] = temp_counts.get(cat, 0)
    
        col_cats = st.columns(len(categorias))
        emojis = ["🌟", "👍", "🟡", "🔴"]
        for i, cat in enumerate(categorias):
            col_cats[i].metric(f"{emojis[i]} {cat.title()}", calif_counts[cat])
    
        st.markdown("---")

    
        st.markdown("<h2 style='font-size:24px;'>🗂️ Distribución por Nivel de Evaluación</h2>", unsafe_allow_html=True)
        df_no_anuladas["formulario"] = df_no_anuladas["formulario"].astype(str)
        niveles_eval = {
            "🔵 Nivel Jerárquico": ["1"],
            "🟣 Niveles Medios": ["2", "3", "4"],
            "🟢 Niveles Operativos": ["5", "6"]
        }
    
        cols = st.columns(3)
        for i, (titulo, formularios) in enumerate(niveles_eval.items()):
            cantidad = df_no_anuladas["formulario"].isin(formularios).sum()
            cols[i].metric(titulo, cantidad)
    
        columnas_necesarias = [
            "cuil", "apellido_nombre", "calificacion", "puntaje_total", "formulario",
            "nivel", "agrupamiento", "ingresante"
        ]
        for col in columnas_necesarias:
            if col not in df_no_anuladas.columns:
                df_no_anuladas[col] = ""
    
    elif seleccion == "✅ EVALUACIONES":

   
  
    
        st.markdown("<br><br>", unsafe_allow_html=True)  # Espacio más grande
    
        # Asegurar columnas necesarias antes de usar la tabla
        for col in ["apellido_nombre", "formulario", "calif_puntaje", "evaluador", "Fecha_formateada"]:
            if col not in df_no_anuladas.columns:
                df_no_anuladas[col] = ""
        
        # ---- TABLA DE EVALUACIONES REGISTRADAS ----
        st.markdown("<h2 style='font-size:20px;'>✅ Evaluaciones registradas:</h2>", unsafe_allow_html=True)
        
        if df_no_anuladas.empty:
            st.info("ℹ️ No hay evaluaciones registradas.")
        else:
            # Crear copia para visualización sin afectar el original
            df_visual = df_no_anuladas.copy()
            
            # Agregar columnas visuales
            df_visual["Nivel Eval"] = df_visual["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
            df_visual["Puntaje/Máximo"] = df_visual.apply(
                lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
                axis=1
            )
            
            # Ordenar para visualización
            df_visual = df_visual.sort_values(by=["apellido_nombre", "Fecha_formateada"])
            
            # Paginación
            registros_por_pagina = 8
            total_registros = len(df_visual)
            total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)
            paginas = list(range(1, total_paginas + 1))
            
            if (
                "pagina_evaluadas" not in st.session_state
                or st.session_state["pagina_evaluadas"] not in paginas
            ):
                st.session_state["pagina_evaluadas"] = 1
            
            pagina_actual = st.selectbox(
                "Seleccionar página:",
                options=paginas,
                index=paginas.index(st.session_state["pagina_evaluadas"]),
                key="pagina_evaluadas_select"
            )
            
            st.session_state["pagina_evaluadas"] = pagina_actual
            
            # Mostrar tabla formateada con paginación
            inicio = (pagina_actual - 1) * registros_por_pagina
            fin = inicio + registros_por_pagina
            
            st.dataframe(
                df_visual.iloc[inicio:fin][[
                    "apellido_nombre", "Nivel Eval", "calificacion",
                    "Puntaje/Máximo", "evaluador", "Fecha_formateada"
                ]].rename(columns={
                    "apellido_nombre": "Apellido y Nombres",
                    "Nivel Eval": "Nivel Evaluación",
                    "calificacion": "Calificación",
                    "Puntaje/Máximo": "Puntaje/Máximo",
                    "evaluador": "Evaluador",
                    "Fecha_formateada": "Fecha"
                }),
                use_container_width=True,
                hide_index=True
            )

            with st.spinner("✏️ Generando documento..."):
                doc = generar_informe_docx(df_informe, df_evaluados, dependencia_filtro)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
        
            with open(tmp_path, "rb") as file:
                st.download_button(
                    label="📥 Descargar Informe",
                    data=file,
                    file_name=f"informe_{dependencia_filtro.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )        
            
        def set_cell_style(cell, bold=True, bg_color=None, font_color="000000"):
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(" ")
            run.text = run.text if run.text.strip() else " "
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(font_color.upper())
        
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            if bg_color:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), bg_color)
                tcPr.append(shd)
        
            para.alignment = 1  # Centrado
        
        
        def generar_informe_docx(df_base, df_eval, dependencia_nombre):
            doc = Document()
        
            # Márgenes 2 cm
            section = doc.sections[0]
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
            # Fuente general
            doc.styles["Normal"].font.name = "Calibri"
            doc.styles["Normal"].font.size = Pt(10)
        
            # Títulos con estilo
            def titulo(texto):
                p = doc.add_paragraph()
                run = p.add_run(texto)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("104f8e")
        
            # Encabezado
            header = section.header
            p_header = header.paragraphs[0]
            p_header.clear()
            run = p_header.add_run(
                f"INSTITUTO NACIONAL DE ESTADISTICA Y CENSOS\n"
                f"DIRECCIÓN DE CAPACITACIÓN Y CARRERA DE PERSONAL\n"
                f"EVALUACIÓN DE DESEMPEÑO 2024\n"
                f"UNIDAD DE ANÁLISIS: {dependencia_nombre}"
            )
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("104f8e")
            p_header.alignment = 1
            p_header.paragraph_format.line_spacing = Pt(12)
    
    
            
            # AGRUPAMIENTO
            doc.add_paragraph()
    
            titulo("PERSONAL TOTAL POR TIPO DE AGRUPAMIENTO")
            gral = len(df_base[df_base["agrupamiento"] == "GRAL"])
            prof = len(df_base[df_base["agrupamiento"] == "PROF"])
            tabla_agrup = doc.add_table(rows=2, cols=2)
            tabla_agrup.style = 'Table Grid'
            for i, h in enumerate(["GENERAL", "PROFESIONAL"]):
                tabla_agrup.cell(0, i).text = h
                set_cell_style(tabla_agrup.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
                tabla_agrup.cell(1, i).text = str([gral, prof][i])
                set_cell_style(tabla_agrup.cell(1, i), bold=False)
        
            doc.add_paragraph()
        
            # ESCALAFÓN
            titulo("PERSONAL TOTAL POR TIPO DE NIVEL ESCALAFONARIO")
            niveles = ["A", "B", "C", "D", "E"]
            conteo_niveles = df_base["nivel"].value_counts()
            tabla_nivel = doc.add_table(rows=2, cols=5)
            tabla_nivel.style = 'Table Grid'
            for i, nivel in enumerate(niveles):
                tabla_nivel.cell(0, i).text = nivel
                set_cell_style(tabla_nivel.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
                tabla_nivel.cell(1, i).text = str(conteo_niveles.get(nivel, 0))
                set_cell_style(tabla_nivel.cell(1, i), bold=False)
        
            doc.add_paragraph()
        
            # EVALUADO / INGRESANTE
            titulo("PERSONAL PARA EVALUAR/EVALUADO")
            df_evaluable = df_base[df_base["ingresante"].isin([True, False])]
            no_ingresantes = len(df_evaluable[df_evaluable["ingresante"] == False])
            ingresantes = len(df_evaluable[df_evaluable["ingresante"] == True])
            total_evaluable = no_ingresantes + ingresantes
            evaluados = df_eval["calificacion"].apply(lambda x: x != "").sum()
        
            tabla_eval = doc.add_table(rows=2, cols=4)
            tabla_eval.style = 'Table Grid'
            eval_headers = [
                "PERMANENTES NO INGRESANTE",
                "PERMANENTES INGRESANTES",
                "TOTAL A EVALUAR",
                "TOTAL EVALUADO"
            ]
            for i, h in enumerate(eval_headers):
                tabla_eval.cell(0, i).text = h
                set_cell_style(tabla_eval.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
                tabla_eval.cell(1, i).text = str([no_ingresantes, ingresantes, total_evaluable, evaluados][i])
                set_cell_style(tabla_eval.cell(1, i), bold=False)
        
            doc.add_paragraph()
        
            # FUNCION AUXILIAR: Listados por formulario
            def agregar_tabla_por_formulario(titulo_texto, formularios):
                titulo(titulo_texto)
                subset = df_eval[df_eval["formulario"].astype(str).isin(formularios)].sort_values("apellido_nombre")
        
                tabla = doc.add_table(rows=1, cols=3)
                tabla.style = 'Table Grid'
                headers = ["APELLIDOS Y NOMBRES", "CALIFICACIÓN", "PUNTAJE"]
                for i, col in enumerate(headers):
                    tabla.cell(0, i).text = col
                    set_cell_style(tabla.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
        
                if subset.empty:
                    doc.add_paragraph("No hay evaluaciones registradas en este nivel.")
                    return
        
                for _, row in subset.iterrows():
                    r = tabla.add_row().cells
                    r[0].text = row.get("apellido_nombre", "")
                    r[1].text = row.get("calificacion", "")
                    puntaje = row.get("puntaje_total", "")
                    try:
                        puntaje_float = float(puntaje)
                        puntaje = int(puntaje_float) if puntaje_float.is_integer() else puntaje_float
                    except:
                        pass
                    r[2].text = str(puntaje)
        
            agregar_tabla_por_formulario("EVALUACIONES - NIVEL JERÁRQUICO (FORMULARIO 1)", ["1"])
            doc.add_paragraph()
            agregar_tabla_por_formulario("EVALUACIONES - NIVELES MEDIO (FORMULARIOS 2, 3 y 4)", ["2", "3", "4"])
            doc.add_paragraph()
            agregar_tabla_por_formulario("EVALUACIONES - NIVELES OPERATIVOS (FORMULARIOS 5 Y 6)", ["5", "6"])
        
            return doc
    
        
       
        df_informe = df_agentes.copy()  # todos los agentes asignados
        
        df_evaluados = df_agentes[["cuil", "apellido_nombre"]].merge(
            df_no_anuladas[["cuil", "formulario", "calificacion", "puntaje_total"]],
            on="cuil", how="left"
        ).fillna("")
    
    
       # st.markdown("---")
       # st.markdown("<hr style='border:2px solid #136ac1;'>", unsafe_allow_html=True) #linea divisora
       # st.markdown("<h3 style='font-size:22px;'>📋 Informe Evaluaciones Realizadas</h3>", unsafe_allow_html=True)


        
        if df_informe.empty:
            st.warning("⚠️ No hay agentes registrados en esta unidad.")
        elif df_no_anuladas.empty:
            st.info("ℹ️ No hay evaluaciones registradas para generar el informe.")
        else:
            for col in ["formulario", "calificacion", "puntaje_total", "apellido_nombre"]:
                if col not in df_evaluados.columns:
                    df_evaluados[col] = ""
        


     
        # Obtener configuración global
        config_items = supabase.table("configuracion").select("*").execute().data
        config_map = {item["id"]: item for item in config_items}
        anulacion_activa = config_map.get("anulacion_activa", {}).get("valor", True)
        
    
        # Mostrar bloque de anulaciones solo si está habilitado
        if anulacion_activa and not df_no_anuladas.empty:
            st.markdown("---")
            st.markdown("<h2 style='font-size:20px;'>🔄 Evaluaciones que pueden anularse:</h2>", unsafe_allow_html=True)
            
            # Columnas auxiliares
            df_no_anuladas["Seleccionar"] = False
            df_no_anuladas["Nivel Eval"] = df_no_anuladas["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
            df_no_anuladas["Puntaje/Max"] = df_no_anuladas.apply(
                lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
                axis=1
            )
            
            # Ordenar
            df_no_anuladas = df_no_anuladas.sort_values(by=["apellido_nombre", "Fecha_formateada"])
            
            # --- Paginación con selectbox (versión segura) ---
            registros_por_pagina = 8
            total_registros = len(df_no_anuladas)
            total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)
            paginas = list(range(1, total_paginas + 1))
            
            if (
                "pagina_anulables" not in st.session_state
                or st.session_state["pagina_anulables"] not in paginas
            ):
                st.session_state["pagina_anulables"] = 1
            
            pagina_actual = st.selectbox(
                "Seleccionar página:",
                options=paginas,
                index=paginas.index(st.session_state["pagina_anulables"]),
                key="pagina_anulables_select"
            )
            
            st.session_state["pagina_anulables"] = pagina_actual
            
            inicio = (pagina_actual - 1) * registros_por_pagina
            fin = inicio + registros_por_pagina
            
            # Subset paginado
            df_pagina = df_no_anuladas.iloc[inicio:fin][[
                "Seleccionar", "apellido_nombre", "Nivel Eval",
                "calificacion", "Puntaje/Max", "evaluador", "id_evaluacion"
            ]].rename(columns={
                "Seleccionar": "Seleccionar",
                "apellido_nombre": "Apellido y Nombres",
                "Nivel Eval": "Nivel Evaluación",
                "calificacion": "Calificación",
                "Puntaje/Max": "Puntaje/Máximo",
                "evaluador": "Evaluador",
                "id_evaluacion": "id_evaluacion"
            })
            
            # Editor con id_evaluacion oculta pero disponible
            seleccion = st.data_editor(
                df_pagina,
                use_container_width=True,
                hide_index=True,
                disabled=[
                    "Apellido y Nombres", "Nivel Evaluación", "Calificación",
                    "Puntaje", "Evaluador", "id_evaluacion"
                ],
                column_config={
                    "Seleccionar": st.column_config.CheckboxColumn("Seleccionar"),
                    "id_evaluacion": None  # ⛔ Oculta visualmente esta columna
                }
            )



            
            # Botón para anular seleccionadas
            if st.button("❌ Anular seleccionadas"):
                if "Seleccionar" in seleccion.columns:
                    seleccionados = seleccion[seleccion["Seleccionar"] == True]
                    ids_seleccionados = seleccionados["id_evaluacion"].tolist()
                else:
                    ids_seleccionados = []
            
                if not ids_seleccionados:
                    st.warning("⚠️ No hay evaluaciones seleccionadas para anular.")
                else:
                    for id_eval in ids_seleccionados:
                        eval_sel = df_no_anuladas[df_no_anuladas["id_evaluacion"] == id_eval].iloc[0]
                        supabase.table("evaluaciones").update({"anulada": True})\
                            .eq("id_evaluacion", id_eval).execute()
                        supabase.table("agentes").update({"evaluado_2024": False})\
                            .eq("cuil", str(eval_sel["cuil"]).strip()).execute()
                    st.success(f"✅ {len(ids_seleccionados)} evaluaciones anuladas.")
                    time.sleep(2)
                    st.rerun()
  #      elif not anulacion_activa:
  #          st.info("🔒 La anulación de evaluaciones está cerrada.")
        # Si la anulación no está activa, no se muestra ningún bloque
        pass

    
        df_anuladas = df_eval[df_eval["anulada"] == True].copy()

        
        if not df_anuladas.empty:
            st.markdown("---")

        
        if not df_anuladas.empty:
            st.markdown("<h2 style='font-size:20px;'>❌ Evaluaciones anuladas:</h2>", unsafe_allow_html=True)
        
            # Crear copia visual
            df_visual_anuladas = df_anuladas.copy()
            df_visual_anuladas["Nivel Eval"] = df_visual_anuladas["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
            df_visual_anuladas["Puntaje/Máximo"] = df_visual_anuladas.apply(
                lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
                axis=1
            )
        
            # Ordenar
            df_visual_anuladas = df_visual_anuladas.sort_values(by=["apellido_nombre", "Fecha_formateada"])
        
            # Configurar paginación
            registros_por_pagina = 8
            total_registros = len(df_visual_anuladas)
            total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)  # Siempre al menos 1 página
            paginas = list(range(1, total_paginas + 1))
            
            # Inicializar página en session_state si no existe o es inválida
            if (
                "pagina_anuladas" not in st.session_state
                or st.session_state["pagina_anuladas"] not in paginas
            ):
                st.session_state["pagina_anuladas"] = 1
            
            # Selector de página con index seguro
            pagina_actual = st.selectbox(
                "Seleccionar página:",
                options=paginas,
                index=paginas.index(st.session_state["pagina_anuladas"]),
                key="pagina_anuladas_select"
            )
            
            # Actualizar session_state
            st.session_state["pagina_anuladas"] = pagina_actual
            
            # Subset paginado
            inicio = (pagina_actual - 1) * registros_por_pagina
            fin = inicio + registros_por_pagina

        
            subset = df_visual_anuladas.iloc[inicio:fin][[
                "apellido_nombre", "Nivel Eval", "calificacion",
                "Puntaje/Máximo", "evaluador", "Fecha_formateada"
            ]].rename(columns={
                "apellido_nombre": "Apellido y Nombres",
                "Nivel Eval": "Nivel Evaluación",
                "calificacion": "Calificación",
                "Puntaje/Máximo": "Puntaje/Máximo",
                "evaluador": "Evaluador",
                "Fecha_formateada": "Fecha",
            })
        
            st.dataframe(subset, use_container_width=True, hide_index=True)
        
        


                
    elif seleccion == "👥 AGENTES":


        st.markdown("<h2 style='font-size:20px;'>Distribución por Nivel Escalafonario</h2>", unsafe_allow_html=True)
        
        niveles = ["A", "B", "C", "D", "E"]
        conteo_niveles = df_agentes["nivel"].value_counts().to_dict()
        total_niveles = sum(conteo_niveles.get(n, 0) for n in niveles)
        
        fig_niv = go.Figure()
        colores_niveles = ['#F1948A', '#F7DC6F', '#82E0AA', '#85C1E9', '#D7BDE2']
        
        for i, nivel in enumerate(niveles):
            cantidad = conteo_niveles.get(nivel, 0)
            pct = cantidad / total_niveles * 100 if total_niveles > 0 else 0
            fig_niv.add_trace(go.Bar(
                y=[""],
                x=[cantidad],
                name=f"NIVEL {nivel}",
                marker_color=colores_niveles[i],
                orientation='h',
                customdata=[[cantidad, pct]],
                hovertemplate=f"NIVEL {nivel}: "+"%{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>"
            ))
        
        # --- Niveles ---
        fig_niv.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_niv, use_container_width=True, config={"displayModeBar": False})


        st.markdown("<h2 style='font-size:20px;'>Distribución por Agrupamiento</h2>", unsafe_allow_html=True)
        
        # Calcular cantidades
        gral = len(df_agentes[df_agentes["agrupamiento"] == "GRAL"])
        prof = len(df_agentes[df_agentes["agrupamiento"] == "PROF"])
        total = prof + gral
        pct_gral = gral / total * 100 if total > 0 else 0
        pct_prof = prof / total * 100 if total > 0 else 0
        
        
        fig_agru = go.Figure()

        fig_agru.add_trace(go.Bar(
            y=[""],
            x=[prof],
            name="PROFESIONAL",
            marker_color='#82E0AA',  # PROF
            orientation='h',
            customdata=[[prof, pct_prof]],
            hovertemplate='PROFESIONAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        
        fig_agru.add_trace(go.Bar(
            y=[""],
            x=[gral],
            name="GENERAL",
            marker_color='#A19AD3',  # GRAL
            orientation='h',
            customdata=[[gral, pct_gral]],
            hovertemplate='GENERAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        

        
        fig_agru.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_agru, use_container_width=True, config={"displayModeBar": False})
        

        st.markdown("<h2 style='font-size:20px;'>Distribución por Tramo</h2>", unsafe_allow_html=True)
        
        # Calcular cantidades
        geral = len(df_agentes[df_agentes["tramo"] == "GENERAL"])
        inter = len(df_agentes[df_agentes["tramo"] == "INTERMEDIO"])
        avanz = len(df_agentes[df_agentes["tramo"] == "AVANZADO"])
        total = geral + inter + avanz
        pct_geral = geral / total * 100 if total > 0 else 0
        pct_inter = inter / total * 100 if total > 0 else 0
        pct_avanz = avanz / total * 100 if total > 0 else 0
        
        
        fig_tram = go.Figure()
        
        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[geral],
            name="GENERAL",
            marker_color='#A19AD3',  # GRAL
            orientation='h',
            customdata=[[geral, pct_geral]],
            hovertemplate='GENERAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[inter],
            name="INTERMEDIO",
            marker_color='#82E0AA',  
            orientation='h',
            customdata=[[inter, pct_inter]],
            hovertemplate='INTERMEDIO: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))

        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[avanz],
            name="AVANZADO",
            marker_color='#F1948A',  
            orientation='h',
            customdata=[[avanz, pct_avanz]],
            hovertemplate='AVANZADO: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
    
        
        fig_tram.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_tram, use_container_width=True, config={"displayModeBar": False})
      


        st.markdown("<h2 style='font-size:20px;'>Distribución por Ingreso en Planta Permanente</h2>", unsafe_allow_html=True)
        
        ingresantes = len(df_agentes[df_agentes["ingresante"] == True])
        no_ingresantes = len(df_agentes[df_agentes["ingresante"] == False])
        total_ing = ingresantes + no_ingresantes
        pct_ing = ingresantes / total_ing * 100 if total_ing > 0 else 0
        pct_no_ing = no_ingresantes / total_ing * 100 if total_ing > 0 else 0
        
        fig_ing = go.Figure()
        
        fig_ing.add_trace(go.Bar(
            y=[""],
            x=[no_ingresantes],
            name="HISTÓRICOS",
            marker_color='#A19AD3',
            orientation='h',
            customdata=[[no_ingresantes, pct_no_ing]],
            hovertemplate='HISTÓRICOS: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_ing.add_trace(go.Bar(
            y=[""],
            x=[ingresantes],
            name="INGRESANTES",
            marker_color='#82E0AA',
            orientation='h',
            customdata=[[ingresantes, pct_ing]],
            hovertemplate='INGRESANTES: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_ing.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
       
        st.plotly_chart(fig_ing, use_container_width=True, config={"displayModeBar": False})
        
