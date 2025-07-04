import streamlit as st
import pandas as pd
import tempfile
import time
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def mostrar_evaluaciones(data):
                 
    elif seleccion == "👥 AGENTES":

        # st.markdown("<h2 style='font-size:20px;'>Agentes evaluables</h2>", unsafe_allow_html=True)
        st.markdown("---")
        cantidad_agentes = len(df_agentes)
        st.markdown(f"<h2 style='font-size:20px;'>👥 Total de agentes evaluables: <strong>{cantidad_agentes}</strong></h2>", unsafe_allow_html=True)


        def set_cell_style(cell, bold=True, bg_color=None, font_color="000000"):
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(" ")
            run.text = run.text if run.text.strip() else " "
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(font_color.upper())

            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            if bg_color:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), bg_color)
                tcPr.append(shd)

            para.alignment = 1  # Centrado



        def generar_informe_agentes_docx(df_agentes, dependencia_nombre):
            doc = Document()

            # Márgenes
            section = doc.sections[0]
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

            # Fuente general
            doc.styles["Normal"].font.name = "Calibri"
            doc.styles["Normal"].font.size = Pt(10)

            # Encabezado
            header = section.header
            p_header = header.paragraphs[0]
            p_header.clear()
            run = p_header.add_run(
                f"INSTITUTO NACIONAL DE ESTADISTICA Y CENSOS\n"
                f"DIRECCIÓN DE CAPACITACIÓN Y CARRERA DE PERSONAL\n"
                f"LISTADO DE AGENTES PARA EVALUACIÓN DE DESEMPEÑO 2024\n"
                f"UNIDAD DE ANÁLISIS: {dependencia_nombre}"
            )
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("104f8e")
            p_header.alignment = 1
            p_header.paragraph_format.line_spacing = Pt(12)

            doc.add_paragraph()

            # Título
            p_title = doc.add_paragraph()
            run = p_title.add_run("LISTADO DE AGENTES")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("104f8e")
            p_title.alignment = 1

            doc.add_paragraph()

            # Tabla
            columnas = ["APELLIDO Y NOMBRE", "NIVEL/GRADO", "AGRUPAMIENTO", "TRAMO", "INGRESANTE"]
            tabla = doc.add_table(rows=1, cols=len(columnas))
            tabla.style = 'Table Grid'

            for i, col in enumerate(columnas):
                celda = tabla.cell(0, i)
                celda.text = col
                set_cell_style(celda, bg_color="104f8e", font_color="FFFFFF")

            # Filtrado y orden
            df_agentes_ordenado = df_agentes.sort_values("apellido_nombre")

            for _, row in df_agentes_ordenado.iterrows():
                fila = tabla.add_row().cells
                fila[0].text = row.get("apellido_nombre", "")
                fila[1].text = f"{row.get('nivel', '')}-{row.get('grado', '')}"
                agrupamiento = row.get("agrupamiento", "")
                fila[2].text = "Profesional" if agrupamiento == "PROF" else "General" if agrupamiento == "GRAL" else ""
                fila[3].text = row.get("tramo", "")
                fila[4].text = "Sí" if row.get("ingresante") is True else ""

                for celda in fila:
                    set_cell_style(celda, bold=False)

            return doc

        # 🔽 Botón de descarga
        if not df_agentes.empty:
            with st.spinner("✏️ Generando documento..."):
                doc = generar_informe_agentes_docx(df_agentes, dependencia_filtro)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name

            with open(tmp_path, "rb") as file:
                st.download_button(
                    label="📥 Descargar Informe de Agentes",
                    data=file,
                    file_name=f"agentes_{dependencia_filtro.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    # use_container_width=True,
                    type="primary"
                )

        st.markdown("<h2 style='font-size:20px;'>Distribución por Nivel Escalafonario</h2>", unsafe_allow_html=True)
        
        niveles = ["A", "B", "C", "D", "E"]
        conteo_niveles = df_agentes["nivel"].value_counts().to_dict()
        total_niveles = sum(conteo_niveles.get(n, 0) for n in niveles)
        
        fig_niv = go.Figure()
        colores_niveles = ['#F1948A', '#F7DC6F', '#82E0AA', '#85C1E9', '#D7BDE2']
        
        for i, nivel in enumerate(niveles):
            cantidad = conteo_niveles.get(nivel, 0)
            pct = cantidad / total_niveles * 100 if total_niveles > 0 else 0
            fig_niv.add_trace(go.Bar(
                y=[""],
                x=[cantidad],
                name=f"NIVEL {nivel}",
                marker_color=colores_niveles[i],
                orientation='h',
                customdata=[[cantidad, pct]],
                hovertemplate=f"NIVEL {nivel}: "+"%{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>"
            ))
        
        # --- Niveles ---
        fig_niv.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_niv, use_container_width=True, config={"displayModeBar": False})

        st.markdown("<h2 style='font-size:20px;'>Distribución por Agrupamiento</h2>", unsafe_allow_html=True)
        
        # Calcular cantidades
        gral = len(df_agentes[df_agentes["agrupamiento"] == "GRAL"])
        prof = len(df_agentes[df_agentes["agrupamiento"] == "PROF"])
        total = prof + gral
        pct_gral = gral / total * 100 if total > 0 else 0
        pct_prof = prof / total * 100 if total > 0 else 0
        
        fig_agru = go.Figure()

        fig_agru.add_trace(go.Bar(
            y=[""],
            x=[prof],
            name="PROFESIONAL",
            marker_color='#82E0AA',  # PROF
            orientation='h',
            customdata=[[prof, pct_prof]],
            hovertemplate='PROFESIONAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_agru.add_trace(go.Bar(
            y=[""],
            x=[gral],
            name="GENERAL",
            marker_color='#A19AD3',  # GRAL
            orientation='h',
            customdata=[[gral, pct_gral]],
            hovertemplate='GENERAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))

        fig_agru.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_agru, use_container_width=True, config={"displayModeBar": False})

        st.markdown("<h2 style='font-size:20px;'>Distribución por Tramo</h2>", unsafe_allow_html=True)
        
        # Calcular cantidades
        geral = len(df_agentes[df_agentes["tramo"] == "GENERAL"])
        inter = len(df_agentes[df_agentes["tramo"] == "INTERMEDIO"])
        avanz = len(df_agentes[df_agentes["tramo"] == "AVANZADO"])
        total = geral + inter + avanz
        pct_geral = geral / total * 100 if total > 0 else 0
        pct_inter = inter / total * 100 if total > 0 else 0
        pct_avanz = avanz / total * 100 if total > 0 else 0
        
        fig_tram = go.Figure()
        
        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[geral],
            name="GENERAL",
            marker_color='#A19AD3',  # GRAL
            orientation='h',
            customdata=[[geral, pct_geral]],
            hovertemplate='GENERAL: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[inter],
            name="INTERMEDIO",
            marker_color='#82E0AA',  
            orientation='h',
            customdata=[[inter, pct_inter]],
            hovertemplate='INTERMEDIO: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))

        fig_tram.add_trace(go.Bar(
            y=[""],
            x=[avanz],
            name="AVANZADO",
            marker_color='#F1948A',  
            orientation='h',
            customdata=[[avanz, pct_avanz]],
            hovertemplate='AVANZADO: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_tram.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
        st.plotly_chart(fig_tram, use_container_width=True, config={"displayModeBar": False})

        st.markdown("<h2 style='font-size:20px;'>Distribución por Ingreso en Planta Permanente</h2>", unsafe_allow_html=True)
        
        ingresantes = len(df_agentes[df_agentes["ingresante"] == True])
        no_ingresantes = len(df_agentes[df_agentes["ingresante"] == False])
        total_ing = ingresantes + no_ingresantes
        pct_ing = ingresantes / total_ing * 100 if total_ing > 0 else 0
        pct_no_ing = no_ingresantes / total_ing * 100 if total_ing > 0 else 0
        
        fig_ing = go.Figure()
        
        fig_ing.add_trace(go.Bar(
            y=[""],
            x=[no_ingresantes],
            name="HISTÓRICOS",
            marker_color='#A19AD3',
            orientation='h',
            customdata=[[no_ingresantes, pct_no_ing]],
            hovertemplate='HISTÓRICOS: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_ing.add_trace(go.Bar(
            y=[""],
            x=[ingresantes],
            name="INGRESANTES",
            marker_color='#82E0AA',
            orientation='h',
            customdata=[[ingresantes, pct_ing]],
            hovertemplate='INGRESANTES: %{customdata[0]} agentes<br>📊 %{customdata[1]:.1f}%<extra></extra>'
        ))
        
        fig_ing.update_layout(
            barmode='stack',
            height=160,
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="top",
                y=-0.6,
                xanchor="center",
                x=0.5,
                font=dict(size=16),
                traceorder='normal'
            ),
            margin=dict(l=30, r=30, t=30, b=30),
            xaxis_title="",
            yaxis_title=""
        )
       
        st.plotly_chart(fig_ing, use_container_width=True, config={"displayModeBar": False})
