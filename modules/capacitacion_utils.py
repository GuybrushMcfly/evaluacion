import pandas as pd
import math
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


def generar_informe_comite_docx(df, unidad_nombre, total, resumen_niveles, path_docx):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    header = sec.header
    p_head = header.paragraphs[0]
    p_head.text = "Evaluación de Desempeño 2024"
    p_head.alignment = 1
    for run in p_head.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_heading("Resumen Evaluaciones", level=1)

    p_unit = doc.add_paragraph()
    run_u = p_unit.add_run(f"Unidad de Evaluación: {unidad_nombre}")
    run_u.bold = True
    run_u.font.name = "Calibri"
    run_u.font.color.rgb = RGBColor(0, 0, 0)

    azul = "B7E0F7"
    grupos = {}

    if "residual" not in df.columns:
        df["residual"] = False
    if "nivel" not in df.columns:
        df["nivel"] = df["formulario"].astype(int)

    # Residuales
    residuales_df = df[df["residual"] == True]
    if not residuales_df.empty:
        grupos["Unidad Residual"] = residuales_df

    # No residuales
    df_nores = df[df["residual"] != True]

    medios_df = df_nores[df_nores["nivel"].isin([2, 3, 4])]
    if not medios_df.empty:
        if len(medios_df) < 6:
            grupos["Niveles Medios"] = medios_df
        else:
            for lvl in [2, 3, 4]:
                tmp = medios_df[medios_df["nivel"] == lvl]
                if not tmp.empty:
                    grupos[f"Nivel {lvl}"] = tmp

    oper_df = df_nores[df_nores["nivel"].isin([5, 6])]
    if not oper_df.empty:
        grupos["Niveles Operativos"] = oper_df

    cols = ["Apellido y Nombre", "CUIL", "Nivel", "Puntaje Absoluto", "Puntaje Relativo", "Calificación"]
    for titulo, tabla_df in grupos.items():
        doc.add_heading(titulo, level=2)

        tbl = doc.add_table(rows=1 + len(tabla_df), cols=len(cols), style="Table Grid")
        for j, c in enumerate(cols):
            cell = tbl.rows[0].cells[j]
            r = cell.paragraphs[0].add_run(c)
            r.bold = True
            r.font.name = "Calibri"
            tc = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), azul)
            tc.append(shd)

        for i, row in enumerate(tabla_df.itertuples(index=False), start=1):
            cells = tbl.rows[i].cells
            cells[0].text = row.apellido_nombre
            cells[1].text = str(row.cuil)
            cells[2].text = str(row.nivel)
            cells[3].text = str(row.puntaje_total)
            cells[4].text = f"{row.puntaje_relativo:.2f}"
            cells[5].text = row.calificacion
            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        # Totales por grupo
        n = len(tabla_df)
        cupo = max(1, math.ceil(n * 0.1))
        tbl_sum = doc.add_table(rows=2, cols=2, style="Table Grid")
        tbl_sum.rows[0].cells[0].text = "Total evaluados"
        tbl_sum.rows[0].cells[1].text = str(n)
        tbl_sum.rows[1].cells[0].text = "BDD correspondientes (10%)"
        tbl_sum.rows[1].cells[1].text = str(cupo)

        for idx in (0, 1):
            c0 = tbl_sum.rows[idx].cells[0]
            tc0 = c0._tc.get_or_add_tcPr()
            sh = OxmlElement('w:shd')
            sh.set(qn('w:val'), 'clear')
            sh.set(qn('w:fill'), azul)
            tc0.append(sh)
            for p in c0.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            for p in tbl_sum.rows[idx].cells[1].paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    doc.add_page_break()

    # Totales generales
    doc.add_heading("Totales Generales", level=2)
    cupo30 = math.floor(total * 0.3)
    cupo10 = max(1, math.ceil(total * 0.1))
    tbl_tot = doc.add_table(rows=3, cols=2, style="Table Grid")
    labels = [
        ("TOTAL DE AGENTES EVALUADOS", str(total)),
        ("CUPO DESTACADOS (30%)", str(cupo30)),
        ("CUPO BONIFICACIÓN ESPECIAL (10%)", str(cupo10)),
    ]
    for idx, (lab, val) in enumerate(labels):
        c0 = tbl_tot.rows[idx].cells[0]
        c1 = tbl_tot.rows[idx].cells[1]
        c0.text = lab
        c1.text = val
        tc0 = c0._tc.get_or_add_tcPr()
        sh0 = OxmlElement('w:shd')
        sh0.set(qn('w:val'), 'clear')
        sh0.set(qn('w:fill'), azul)
        tc0.append(sh0)
        for p in c0.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(9)
        for p in c1.paragraphs:
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)

    doc.save(path_docx)


def generar_anexo_iii_docx(texto, path_docx):
    doc = Document()
    doc.add_heading("ANEXO III - ACTA DE VEEDURÍA GREMIAL", level=1)
    doc.add_paragraph(texto.strip())
    doc.save(path_docx)


def generar_cuadro_resumen_docx(df_resumen, path_docx):
    doc = Document()
    doc.add_heading("Cuadro Resumen de Niveles", level=1)
    niveles = list(df_resumen.columns)
    filas = list(df_resumen.index)
    table = doc.add_table(rows=len(filas) + 1, cols=len(niveles) + 1)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Nivel"
    for j, nivel in enumerate(niveles, start=1):
        table.rows[0].cells[j].text = str(nivel)
    for i, fila in enumerate(filas, start=1):
        cells = table.rows[i].cells
        cells[0].text = fila
        for j, nivel in enumerate(niveles, start=1):
            cells[j].text = str(df_resumen.loc[fila, nivel])
    doc.save(path_docx)


def analizar_evaluaciones_residuales(df):
    df = df.copy()
    df["nivel"] = df["formulario"].astype(int)
    df["residual"] = False
    df.loc[df["nivel"] == 1, "residual"] = True
    for ua, df_ua in df.groupby("unidad_analisis"):
        medios = df_ua[df_ua["nivel"].isin([2, 3, 4])]
        if len(medios) < 6 and len(medios) > 0:
            df.loc[medios.index, "residual"] = True
        elif len(medios) >= 6:
            grupos_chicos = []
            for n in [2, 3, 4]:
                nivel_n = medios[medios["nivel"] == n]
                if 0 < len(nivel_n) < 6:
                    grupos_chicos.append(nivel_n)
            if grupos_chicos:
                grupo_unificado = pd.concat(grupos_chicos)
                if len(grupo_unificado) < 6:
                    df.loc[grupo_unificado.index, "residual"] = True
                else:
                    df.loc[grupo_unificado.index, "residual"] = False

        operativos = df_ua[df_ua["nivel"].isin([5, 6])]
        if len(operativos) < 6 and len(operativos) > 0:
            df.loc[operativos.index, "residual"] = True
        elif len(operativos) >= 6:
            grupos_chicos = []
            for n in [5, 6]:
                nivel_n = operativos[operativos["nivel"] == n]
                if 0 < len(nivel_n) < 6:
                    grupos_chicos.append(nivel_n)
            if grupos_chicos:
                grupo_unificado = pd.concat(grupos_chicos)
                if len(grupo_unificado) < 6:
                    df.loc[grupo_unificado.index, "residual"] = True
                else:
                    df.loc[grupo_unificado.index, "residual"] = False
    return df
