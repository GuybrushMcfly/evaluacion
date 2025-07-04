# import streamlit as st
import pandas as pd
import tempfile
import time
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def mostrar_evaluaciones(data):

    st.markdown("<br><br>", unsafe_allow_html=True)  # Espacio más grande

    # Asegurar columnas necesarias antes de usar la tabla
    for col in ["apellido_nombre", "formulario", "calif_puntaje", "evaluador", "Fecha_formateada"]:
        if col not in df_no_anuladas.columns:
            df_no_anuladas[col] = ""
    
    # ---- TABLA DE EVALUACIONES REGISTRADAS ----
    st.markdown("<h2 style='font-size:20px;'>✅ Evaluaciones registradas:</h2>", unsafe_allow_html=True)
    
    if df_no_anuladas.empty:
        st.info("ℹ️ No hay evaluaciones registradas.")
    else:
        # Crear copia para visualización sin afectar el original
        df_visual = df_no_anuladas.copy()
        
        # Agregar columnas visuales
        df_visual["Nivel Eval"] = df_visual["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
        df_visual["Puntaje/Máximo"] = df_visual.apply(
            lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
            axis=1
        )
        
        # Ordenar para visualización
        df_visual = df_visual.sort_values(by=["apellido_nombre", "Fecha_formateada"])
        
        # Paginación
        registros_por_pagina = 8
        total_registros = len(df_visual)
        total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)
        paginas = list(range(1, total_paginas + 1))
        
        if (
            "pagina_evaluadas" not in st.session_state
            or st.session_state["pagina_evaluadas"] not in paginas
        ):
            st.session_state["pagina_evaluadas"] = 1
        
        pagina_actual = st.selectbox(
            "Seleccionar página:",
            options=paginas,
            index=paginas.index(st.session_state["pagina_evaluadas"]),
            key="pagina_evaluadas_select"
        )
        
        st.session_state["pagina_evaluadas"] = pagina_actual
        
        # Mostrar tabla formateada con paginación
        inicio = (pagina_actual - 1) * registros_por_pagina
        fin = inicio + registros_por_pagina
        
        st.dataframe(
            df_visual.iloc[inicio:fin][[
                "apellido_nombre", "Nivel Eval", "calificacion",
                "Puntaje/Máximo", "evaluador", "Fecha_formateada"
            ]].rename(columns={
                "apellido_nombre": "Apellido/s y Nombre/s",
                "Nivel Eval": "Nivel Evaluación",
                "calificacion": "Calificación",
                "Puntaje/Máximo": "Puntaje/Máximo",
                "evaluador": "Evaluador",
                "Fecha_formateada": "Fecha"
            }),
            use_container_width=True,
            hide_index=True
        )

    def set_cell_style(cell, bold=True, bg_color=None, font_color="000000"):
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(" ")
        run.text = run.text if run.text.strip() else " "
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(font_color.upper())
    
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        if bg_color:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg_color)
            tcPr.append(shd)
    
        para.alignment = 1  # Centrado

    def generar_informe_docx(df_base, df_eval, dependencia_nombre):
        doc = Document()
    
        # Márgenes 2 cm
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
        # Fuente general
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)
    
        # Títulos con estilo
        def titulo(texto):
            p = doc.add_paragraph()
            run = p.add_run(texto)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("104f8e")
    
        # Encabezado
        header = section.header
        p_header = header.paragraphs[0]
        p_header.clear()
        run = p_header.add_run(
            f"INSTITUTO NACIONAL DE ESTADISTICA Y CENSOS\n"
            f"DIRECCIÓN DE CAPACITACIÓN Y CARRERA DE PERSONAL\n"
            f"EVALUACIÓN DE DESEMPEÑO 2024\n"
            f"UNIDAD DE ANÁLISIS: {dependencia_nombre}"
        )
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("104f8e")
        p_header.alignment = 1
        p_header.paragraph_format.line_spacing = Pt(12)

        # AGRUPAMIENTO
        doc.add_paragraph()

        titulo("PERSONAL TOTAL POR TIPO DE AGRUPAMIENTO")
        gral = len(df_base[df_base["agrupamiento"] == "GRAL"])
        prof = len(df_base[df_base["agrupamiento"] == "PROF"])
        tabla_agrup = doc.add_table(rows=2, cols=2)
        tabla_agrup.style = 'Table Grid'
        for i, h in enumerate(["GENERAL", "PROFESIONAL"]):
            tabla_agrup.cell(0, i).text = h
            set_cell_style(tabla_agrup.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
            tabla_agrup.cell(1, i).text = str([gral, prof][i])
            set_cell_style(tabla_agrup.cell(1, i), bold=False)
    
        doc.add_paragraph()
    
        # ESCALAFÓN
        titulo("PERSONAL TOTAL POR TIPO DE NIVEL ESCALAFONARIO")
        niveles = ["A", "B", "C", "D", "E"]
        conteo_niveles = df_base["nivel"].value_counts()
        tabla_nivel = doc.add_table(rows=2, cols=5)
        tabla_nivel.style = 'Table Grid'
        for i, nivel in enumerate(niveles):
            tabla_nivel.cell(0, i).text = nivel
            set_cell_style(tabla_nivel.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
            tabla_nivel.cell(1, i).text = str(conteo_niveles.get(nivel, 0))
            set_cell_style(tabla_nivel.cell(1, i), bold=False)
    
        doc.add_paragraph()
    
        # EVALUADO / INGRESANTE
        titulo("PERSONAL PARA EVALUAR/EVALUADO")
        df_evaluable = df_base[df_base["ingresante"].isin([True, False])]
        no_ingresantes = len(df_evaluable[df_evaluable["ingresante"] == False])
        ingresantes = len(df_evaluable[df_evaluable["ingresante"] == True])
        total_evaluable = no_ingresantes + ingresantes
        evaluados = df_eval["calificacion"].apply(lambda x: x != "").sum()
    
        tabla_eval = doc.add_table(rows=2, cols=4)
        tabla_eval.style = 'Table Grid'
        eval_headers = [
            "PERMANENTES NO INGRESANTE",
            "PERMANENTES INGRESANTES",
            "TOTAL A EVALUAR",
            "TOTAL EVALUADO"
        ]
        for i, h in enumerate(eval_headers):
            tabla_eval.cell(0, i).text = h
            set_cell_style(tabla_eval.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
            tabla_eval.cell(1, i).text = str([no_ingresantes, ingresantes, total_evaluable, evaluados][i])
            set_cell_style(tabla_eval.cell(1, i), bold=False)
    
        doc.add_paragraph()
    
        # FUNCION AUXILIAR: Listados por formulario
        def agregar_tabla_por_formulario(titulo_texto, formularios):
            titulo(titulo_texto)
            subset = df_eval[df_eval["formulario"].astype(str).isin(formularios)].sort_values("apellido_nombre")
    
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'
            headers = ["APELLIDOS Y NOMBRES", "CALIFICACIÓN", "PUNTAJE"]
            for i, col in enumerate(headers):
                tabla.cell(0, i).text = col
                set_cell_style(tabla.cell(0, i), bg_color="104f8e", font_color="FFFFFF")
    
            if subset.empty:
                doc.add_paragraph("No hay evaluaciones registradas en este nivel.")
                return
    
            for _, row in subset.iterrows():
                r = tabla.add_row().cells
                r[0].text = row.get("apellido_nombre", "")
                r[1].text = row.get("calificacion", "")
                puntaje = row.get("puntaje_total", "")
                try:
                    puntaje_float = float(puntaje)
                    puntaje = int(puntaje_float) if puntaje_float.is_integer() else puntaje_float
                except:
                    pass
                r[2].text = str(puntaje)
    
        agregar_tabla_por_formulario("EVALUACIONES - NIVEL JERÁRQUICO (FORMULARIO 1)", ["1"])
        doc.add_paragraph()
        agregar_tabla_por_formulario("EVALUACIONES - NIVELES MEDIO (FORMULARIOS 2, 3 y 4)", ["2", "3", "4"])
        doc.add_paragraph()
        agregar_tabla_por_formulario("EVALUACIONES - NIVELES OPERATIVOS (FORMULARIOS 5 Y 6)", ["5", "6"])
    
        return doc

    df_informe = df_agentes.copy()  # todos los agentes asignados
    
    df_evaluados = df_agentes[["cuil", "apellido_nombre"]].merge(
        df_no_anuladas[["cuil", "formulario", "calificacion", "puntaje_total"]],
        on="cuil", how="left"
    ).fillna("")

    # st.markdown("---")
    # st.markdown("<hr style='border:2px solid #136ac1;'>", unsafe_allow_html=True) #linea divisora
    # st.markdown("<h3 style='font-size:22px;'>📋 Informe Evaluaciones Realizadas</h3>", unsafe_allow_html=True)

    if df_informe.empty:
        st.warning("⚠️ No hay agentes registrados en esta unidad.")
    elif df_no_anuladas.empty:
        st.info("ℹ️ No hay evaluaciones registradas para generar el informe.")
    else:
        for col in ["formulario", "calificacion", "puntaje_total", "apellido_nombre"]:
            if col not in df_evaluados.columns:
                df_evaluados[col] = ""

        with st.spinner("✏️ Generando documento..."):
            doc = generar_informe_docx(df_informe, df_evaluados, dependencia_filtro)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name
    
        with open(tmp_path, "rb") as file:
            st.download_button(
                label="📥 Descargar Informe",
                data=file,
                file_name=f"informe_{dependencia_filtro.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                # use_container_width=True,
                type="primary"
            )

    # Mostrar tabla solo para evaluador_general con destacados y si hay al menos 4 agentes activos
    if tiene_rol("evaluador_general") and not df_no_anuladas.empty:
        dependencia_actual = df_agentes["dependencia_general"].dropna().unique()
        dependencia_actual = dependencia_actual[0] if len(dependencia_actual) > 0 else ""
    
        if dependencia_actual:
            try:
                # Obtener total de agentes activos en esa dependencia_general
                agentes_dependencia = supabase.table("agentes")\
                    .select("cuil, activo")\
                    .eq("dependencia_general", dependencia_actual)\
                    .execute().data

                total_activos = sum(1 for a in agentes_dependencia if a.get("activo") == True)
                
                if total_activos > 3:  # Mostrar solo si hay más de 3
                    import math
                    cupo_exacto = total_activos * 0.3
                    max_destacados = math.floor(cupo_exacto) if (cupo_exacto - math.floor(cupo_exacto)) <= 0.5 else math.floor(cupo_exacto) + 1
                
                    df_destacados = df_no_anuladas[df_no_anuladas["calificacion"] == "DESTACADO"].copy()
                    usados = len(df_destacados)
                
                    st.markdown("---")
                    st.markdown(f"<h2 style='font-size:20px;'>🌟 Evaluaciones con calificación DESTACADO ({usados} / {max_destacados})</h2>", unsafe_allow_html=True)
    
                    df_destacados["Nivel Eval"] = df_destacados["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
                    df_destacados["Puntaje/Máximo"] = df_destacados.apply(
                        lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
                        axis=1
                    )
    
                    df_destacados = df_destacados.sort_values(by=["apellido_nombre", "Fecha_formateada"])
    
                    st.dataframe(
                        df_destacados[[
                            "apellido_nombre", "Nivel Eval", "calificacion",
                            "Puntaje/Máximo", "evaluador", "Fecha_formateada"
                        ]].rename(columns={
                            "apellido_nombre": "Apellido/s y Nombre/s",
                            "Nivel Eval": "Nivel Evaluación",
                            "calificacion": "Calificación",
                            "Puntaje/Máximo": "Puntaje/Máximo",
                            "evaluador": "Evaluador",
                            "Fecha_formateada": "Fecha"
                        }),
                        use_container_width=True,
                        hide_index=True
                    )
            except Exception as e:
                st.warning(f"⚠️ Error al calcular destacados disponibles: {e}")

    # Obtener configuración global
    config_items = supabase.table("configuracion").select("*").execute().data
    config_map = {item["id"]: item for item in config_items}
    anulacion_activa = config_map.get("anulacion_activa", {}).get("valor", True)

    # Mostrar bloque de anulaciones solo si está habilitado
    if anulacion_activa and not df_no_anuladas.empty:
        st.markdown("---")
        st.markdown("<h2 style='font-size:20px;'>🔄 Evaluaciones que pueden anularse:</h2>", unsafe_allow_html=True)
        
        # Columnas auxiliares
        df_no_anuladas["Seleccionar"] = False
        df_no_anuladas["Nivel Eval"] = df_no_anuladas["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
        df_no_anuladas["Puntaje/Max"] = df_no_anuladas.apply(
            lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
            axis=1
        )
        
        # Ordenar
        df_no_anuladas = df_no_anuladas.sort_values(by=["apellido_nombre", "Fecha_formateada"])
        
        # --- Paginación con selectbox (versión segura) ---
        registros_por_pagina = 8
        total_registros = len(df_no_anuladas)
        total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)
        paginas = list(range(1, total_paginas + 1))
        
        if (
            "pagina_anulables" not in st.session_state
            or st.session_state["pagina_anulables"] not in paginas
        ):
            st.session_state["pagina_anulables"] = 1
        
        pagina_actual = st.selectbox(
            "Seleccionar página:",
            options=paginas,
            index=paginas.index(st.session_state["pagina_anulables"]),
            key="pagina_anulables_select"
        )
        
        st.session_state["pagina_anulables"] = pagina_actual
        
        inicio = (pagina_actual - 1) * registros_por_pagina
        fin = inicio + registros_por_pagina
        
        # Subset paginado
        df_pagina = df_no_anuladas.iloc[inicio:fin][[
            "Seleccionar", "apellido_nombre", "Nivel Eval",
            "calificacion", "Puntaje/Max", "evaluador", "id_evaluacion"
        ]].rename(columns={
            "Seleccionar": "Seleccionar",
            "apellido_nombre": "Apellido/s y Nombre/s",
            "Nivel Eval": "Nivel Evaluación",
            "calificacion": "Calificación",
            "Puntaje/Max": "Puntaje/Máximo",
            "evaluador": "Evaluador",
            "id_evaluacion": "id_evaluacion"
        })
        
        # Editor con id_evaluacion oculta pero disponible
        seleccion = st.data_editor(
            df_pagina,
            use_container_width=True,
            hide_index=True,
            disabled=[
                "Apellido/s y Nombre/s", "Nivel Evaluación", "Calificación",
                "Puntaje", "Evaluador", "id_evaluacion"
            ],
            column_config={
                "Seleccionar": st.column_config.CheckboxColumn("Seleccionar"),
                "id_evaluacion": None  # ⛔ Oculta visualmente esta columna
            }
        )

        # Botón para anular seleccionadas
        # if st.button("❌ Anular seleccionadas", use_container_width=True, type="primary"):
        if st.button(
            "❌ Anular seleccionadas",
            # use_container_width=True,
            type="primary",
            help="⚠️ Esto eliminará las entradas seleccionadas. No se puede deshacer.",  # Tooltip
        ):
            if "Seleccionar" in seleccion.columns:
                seleccionados = seleccion[seleccion["Seleccionar"] == True]
                ids_seleccionados = seleccionados["id_evaluacion"].tolist()
            else:
                ids_seleccionados = []
        
            if not ids_seleccionados:
                st.warning("⚠️ No hay evaluaciones seleccionadas para anular.")
            else:
                for id_eval in ids_seleccionados:
                    eval_sel = df_no_anuladas[df_no_anuladas["id_evaluacion"] == id_eval].iloc[0]
                    supabase.table("evaluaciones").update({"anulada": True})\
                        .eq("id_evaluacion", id_eval).execute()
                    supabase.table("agentes").update({"evaluado_2024": False})\
                        .eq("cuil", str(eval_sel["cuil"]).strip()).execute()
                st.success(f"✅ {len(ids_seleccionados)} evaluaciones anuladas.")
                time.sleep(2)
                st.rerun()
    # elif not anulacion_activa:
    #     st.info("🔒 La anulación de evaluaciones está cerrada.")
    # Si la anulación no está activa, no se muestra ningún bloque
    pass

    df_anuladas = df_eval[df_eval["anulada"] == True].copy()

    if not df_anuladas.empty:
        st.markdown("---")

    if not df_anuladas.empty:
        st.markdown("<h2 style='font-size:20px;'>❌ Evaluaciones anuladas:</h2>", unsafe_allow_html=True)
    
        # Crear copia visual
        df_visual_anuladas = df_anuladas.copy()
        df_visual_anuladas["Nivel Eval"] = df_visual_anuladas["formulario"].astype(str).map(MAPA_NIVEL_EVALUACION)
        df_visual_anuladas["Puntaje/Máximo"] = df_visual_anuladas.apply(
            lambda row: f"{row['puntaje_total']}/{MAXIMO_PUNTAJE_FORMULARIO.get(str(row['formulario']), '-')}",
            axis=1
        )
    
        # Ordenar
        df_visual_anuladas = df_visual_anuladas.sort_values(by=["apellido_nombre", "Fecha_formateada"])
    
        # Configurar paginación
        registros_por_pagina = 8
        total_registros = len(df_visual_anuladas)
        total_paginas = max(1, (total_registros - 1) // registros_por_pagina + 1)  # Siempre al menos 1 página
        paginas = list(range(1, total_paginas + 1))
        
        # Inicializar página en session_state si no existe o es inválida
        if (
            "pagina_anuladas" not in st.session_state
            or st.session_state["pagina_anuladas"] not in paginas
        ):
            st.session_state["pagina_anuladas"] = 1
        
        # Selector de página con index seguro
        pagina_actual = st.selectbox(
            "Seleccionar página:",
            options=paginas,
            index=paginas.index(st.session_state["pagina_anuladas"]),
            key="pagina_anuladas_select"
        )
        
        # Actualizar session_state
        st.session_state["pagina_anuladas"] = pagina_actual
        
        # Subset paginado
        inicio = (pagina_actual - 1) * registros_por_pagina
        fin = inicio + registros_por_pagina

        subset = df_visual_anuladas.iloc[inicio:fin][[
            "apellido_nombre", "Nivel Eval", "calificacion",
            "Puntaje/Máximo", "evaluador", "Fecha_formateada"
        ]].rename(columns={
            "apellido_nombre": "Apellido/s y Nombre/s",
            "Nivel Eval": "Nivel Evaluación",
            "calificacion": "Calificación",
            "Puntaje/Máximo": "Puntaje/Máximo",
            "evaluador": "Evaluador",
            "Fecha_formateada": "Fecha",
        })
    
        st.dataframe(subset, use_container_width=True, hide_index=True)
